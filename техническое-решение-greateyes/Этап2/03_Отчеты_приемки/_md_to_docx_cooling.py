# One-off / regenerate: python _md_to_docx_cooling.py
# Requires: pip install python-docx (project venv OK)

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor


def _ensure_styles(doc: Document) -> None:
    styles = doc.styles
    try:
        _ = styles["Body Text"]
    except KeyError:
        pass
    # Normal: readable defaults
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.15


def _add_horizontal_rule(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(12)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "BFBFBF")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_runs_with_formatting(p, text: str) -> None:
    """Inline **bold**, `code`, *italic* (single asterisk pair)."""
    if not text:
        return
    pattern = re.compile(r"(\*\*.+?\*\*|`[^`]+`|\*[^*]+\*)")

    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            p.add_run(text[pos : m.start()])
        tok = m.group(1)
        if tok.startswith("**") and tok.endswith("**"):
            r = p.add_run(tok[2:-2])
            r.bold = True
        elif tok.startswith("`") and tok.endswith("`"):
            r = p.add_run(tok[1:-1])
            r.font.name = "Consolas"
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        elif tok.startswith("*") and tok.endswith("*"):
            r = p.add_run(tok[1:-1])
            r.italic = True
        else:
            p.add_run(tok)
        pos = m.end()
    if pos < len(text):
        p.add_run(text[pos:])


def _is_table_separator(line: str) -> bool:
    s = line.strip()
    if not s.startswith("|"):
        return False
    inner = s.strip("|").replace(" ", "")
    return bool(re.fullmatch(r"[\|\-:]+", inner.replace("|", ""))) and "-" in inner


def _table_row_cells(line: str) -> list[str]:
    row = [c.strip() for c in line.strip().split("|")]
    if row and row[0] == "":
        row = row[1:]
    if row and row[-1] == "":
        row = row[:-1]
    return row


def _style_table(table) -> None:
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    run.font.size = Pt(10)
    # Header row bold
    if table.rows:
        for cell in table.rows[0].cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.bold = True
                if not p.runs:
                    r = p.add_run(cell.text)
                    r.bold = True


def convert_md_to_docx(md_path: Path, docx_path: Path) -> None:
    lines = md_path.read_text(encoding="utf-8").splitlines()
    doc = Document()
    _ensure_styles(doc)

    # Title page feel: first # is document title
    i = 0
    n = len(lines)

    while i < n:
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped == "---":
            _add_horizontal_rule(doc)
            i += 1
            continue

        if stripped.startswith("# ") and not stripped.startswith("##"):
            doc.add_heading(stripped[2:].strip(), level=0)
            i += 1
            continue

        if stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=3)
            i += 1
            continue

        if stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=1)
            i += 1
            continue

        # Table block
        if stripped.startswith("|") and "|" in stripped[1:]:
            rows: list[list[str]] = []
            while i < n and lines[i].strip().startswith("|"):
                row_line = lines[i].strip()
                if _is_table_separator(row_line):
                    i += 1
                    continue
                rows.append(_table_row_cells(row_line))
                i += 1
            if rows:
                ncol = max(len(r) for r in rows)
                for r in rows:
                    while len(r) < ncol:
                        r.append("")
                table = doc.add_table(rows=len(rows), cols=ncol)
                table.autofit = True
                for ri, row_cells in enumerate(rows):
                    for ci, val in enumerate(row_cells):
                        cell = table.rows[ri].cells[ci]
                        cell.text = ""
                        p = cell.paragraphs[0]
                        _add_runs_with_formatting(p, val)
                _style_table(table)
                doc.add_paragraph()
            continue

        # Numbered list
        num_m = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if num_m:
            try:
                p = doc.add_paragraph(style="List Number")
            except KeyError:
                p = doc.add_paragraph(style="List Paragraph")
            _add_runs_with_formatting(p, num_m.group(2))
            i += 1
            continue

        # Bullet (nested by leading spaces in original line)
        bullet_m = re.match(r"^(\s*)-\s+(.*)$", line)
        if bullet_m:
            indent = len(bullet_m.group(1).replace("\t", "    "))
            level = min(indent // 2, 3)
            try:
                p = doc.add_paragraph(style="List Bullet")
            except KeyError:
                p = doc.add_paragraph(style="List Paragraph")
            p.paragraph_format.left_indent = Cm(0.5 * (level + 1))
            _add_runs_with_formatting(p, bullet_m.group(2))
            i += 1
            continue

        # Normal paragraph
        p = doc.add_paragraph()
        _add_runs_with_formatting(p, stripped)
        i += 1

    # Page margins (A4 friendly)
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2)

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(docx_path))


def main() -> None:
    base = Path(__file__).resolve().parent
    md = base / "Устранение_багов_по_охлаждению_и_нагреву.md"
    out = base / "Устранение_багов_по_охлаждению_и_нагреву.docx"
    if len(sys.argv) >= 2:
        md = Path(sys.argv[1])
    if len(sys.argv) >= 3:
        out = Path(sys.argv[2])
    if not md.is_file():
        print("Missing MD:", md, file=sys.stderr)
        sys.exit(1)
    convert_md_to_docx(md, out)
    print("Wrote", out)


if __name__ == "__main__":
    main()
